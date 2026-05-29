from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, BackgroundTasks, Request
from services.document_processor import DocumentProcessor
from api.routes.auth import get_authenticated_user
from api.routes.stats import log_activity
from models.auth_models import User
import os
import shutil
import json
import io
import logging
from datetime import datetime
from config import settings
import traceback
from services.thumbnail_service import thumbnail_service
from fastapi.responses import FileResponse, StreamingResponse
import docx2txt
import fitz  # PyMuPDF
from functools import lru_cache
from pydantic import BaseModel
from typing import Optional, List

logger = logging.getLogger(__name__)
router = APIRouter()
doc_processor = DocumentProcessor()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_file_path(user_id: str, course_id: str, filename: str) -> str:
    from utils.file_utils import get_user_file_path
    return get_user_file_path(user_id, course_id, filename)


def _get_annotations_path(user_id: str, course_id: str, filename: str) -> str:
    from utils.file_utils import get_user_annotations_path
    return get_user_annotations_path(user_id, course_id, filename)


def _load_annotations(path: str) -> list:
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _save_annotations(path: str, annotations: list):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(annotations, f, ensure_ascii=False, indent=2)


def _extract_paragraphs(file_path: str) -> List[dict]:
    """
    Extract document content as structured blocks preserving formatting context.
    Returns list of dicts: {text, style} where style is 'heading1','heading2','body','code','list'
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
        blocks = []
        for p in raw.split("\n\n"):
            p = p.strip()
            if p:
                blocks.append({"text": p, "style": "body"})
        return blocks

    elif ext == ".docx":
        import docx
        doc = docx.Document(file_path)
        blocks = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style_name = (p.style.name or "").lower()
            if "heading 1" in style_name:
                style = "heading1"
            elif "heading 2" in style_name:
                style = "heading2"
            elif "heading" in style_name:
                style = "heading3"
            elif p.style.name in ("List Bullet", "List Number", "List Paragraph"):
                style = "list"
            else:
                style = "body"
            blocks.append({"text": text, "style": style})
        return blocks

    elif ext == ".pdf":
        try:
            doc = fitz.open(file_path)
            blocks = []
            for page_num, page in enumerate(doc, 1):
                blocks.append({"text": f"Page {page_num}", "style": "page_break"})
                text_blocks = page.get_text("blocks")
                for b in text_blocks:
                    block_text = b[4].strip()
                    if not block_text:
                        continue
                    for line in block_text.split("\n"):
                        line = line.strip()
                        if line:
                            blocks.append({"text": line, "style": _classify_pdf_line(line)})
            doc.close()
            return blocks
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed, falling back to PyPDF: {e}")
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            blocks = []
            for page_num, page in enumerate(docs, 1):
                blocks.append({"text": f"Page {page_num}", "style": "page_break"})
                for chunk in page.page_content.split("\n\n"):
                    chunk = chunk.strip()
                    if chunk:
                        blocks.append({"text": chunk, "style": "body"})
            return blocks
    elif ext in [".xlsx", ".xls", ".csv", ".xml", ".jpg", ".jpeg", ".png"]:
        # For these types, we rely on the full text extracted during processing
        # and stored in the vector store, but for the 'paragraphs' view, 
        # we can just return the raw extraction or a placeholder.
        try:
            return [{"text": doc_processor.extract_text(file_path), "style": "body"}]
        except:
            return [{"text": f"[Content of {os.path.basename(file_path)}]", "style": "body"}]
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _classify_pdf_line(text: str) -> str:
    """Heuristic to classify a PDF paragraph's visual style."""
    stripped = text.strip()
    # Short ALL-CAPS or ends with no period → likely a heading
    if len(stripped) < 80 and stripped.isupper():
        return "heading1"
    # Numbered section like "1.2 Something" or "Chapter 3"
    import re
    if re.match(r'^(\d+\.)+\d*\s+\S', stripped) or re.match(r'^(Chapter|Section|CHAPTER|SECTION)\s+\d', stripped):
        return "heading2"
    # Bullet / list item
    if re.match(r'^[\u2022\u2023\u25e6\-\*]\s', stripped) or re.match(r'^\d+\.\s', stripped):
        return "list"
    return "body"


ANNOTATION_LABELS = {
    "note": "📝 Note",
    "summary": "📖 Summary",
    "explanation": "💡 Explanation",
    "question": "❓ Question",
    "key_point": "⭐ Key Point",
}


def _build_annotated_docx(paragraphs: List[str], annotations: list, original_filename: str) -> bytes:
    """
    Build a Word document that interleaves original paragraphs with
    inline annotation blocks styled distinctly.
    """
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = docx.Document()

    # Title
    title = doc.add_heading(f"Annotated: {original_filename}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta.runs[0].font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # Build a lookup: paragraph_index -> list of annotations
    ann_map: dict = {}
    for ann in annotations:
        idx = ann.get("paragraph_index", -1)
        ann_map.setdefault(idx, []).append(ann)

    def add_annotation_block(ann):
        label = ANNOTATION_LABELS.get(ann.get("type", "note"), "📝 Note")
        ts = ann.get("created_at", "")[:16].replace("T", " ")

        # Shaded paragraph for the annotation
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

        # Add shading via XML
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "FFF3CD")  # soft yellow
        pPr.append(shd)

        # Label run
        label_run = p.add_run(f"{label}  ")
        label_run.bold = True
        label_run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
        label_run.font.size = Pt(9)

        # Timestamp run
        ts_run = p.add_run(f"({ts})\n")
        ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        ts_run.font.size = Pt(8)
        ts_run.italic = True

        # Content run
        content_run = p.add_run(ann["content"])
        content_run.font.size = Pt(10)
        content_run.font.color.rgb = RGBColor(0x33, 0x33, 0x00)

    # Render paragraphs + inline annotations
    for i, para_text in enumerate(paragraphs):
        # Original paragraph
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(6)

        # Annotations attached to this paragraph
        for ann in ann_map.get(i, []):
            add_annotation_block(ann)

    # Annotations not tied to any paragraph (paragraph_index == -1)
    orphans = ann_map.get(-1, [])
    if orphans:
        doc.add_heading("General Annotations", level=2)
        for ann in orphans:
            add_annotation_block(ann)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@router.post("/upload")
async def upload_document(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    course_id: str = Form(...),
    current_user: User = Depends(get_authenticated_user)
):
    """Upload and process document (fast response, heavy lifting in background)"""
    try:
        user_id = str(current_user.id)
        allowed_extensions = ['.pdf', '.txt', '.docx', '.xlsx', '.xls', '.csv', '.xml', '.jpg', '.jpeg', '.png']
        file_ext = os.path.splitext(file.filename)[1].lower()

        if file_ext not in allowed_extensions:
            raise HTTPException(400, f"Unsupported file type: {file_ext}")

        from utils.file_utils import get_user_course_dir, get_user_file_path, sanitize_filename
        
        original_filename = file.filename
        sanitized_filename = sanitize_filename(original_filename)
        
        user_dir = get_user_course_dir(user_id, course_id)
        os.makedirs(user_dir, exist_ok=True)
        file_path = get_user_file_path(user_id, course_id, sanitized_filename)

        # 1. Save file (very fast)
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 2. Log basic activity
        log_activity(user_id, "document_upload_start", {
            "filename": original_filename,
            "course": course_id
        })

        # 3. Schedule heavy processing in background
        # Pass the API key with multiple fallbacks: Header > State > Database
        api_key = request.headers.get("X-Groq-API-Key") or getattr(request.state, "groq_api_key", None)
        
        if not api_key:
            from services.auth_service import auth_service
            api_key = auth_service.get_groq_key(user_id)
            if api_key:
                logger.info(f"Retrieved Groq key from database for user {user_id}")
        
        background_tasks.add_task(
            doc_processor.process_document, 
            file_path, user_id, course_id, sanitized_filename,
            api_key=api_key
        )

        return {
            "message": "Upload successful! We are processing your document in the background. It will be ready for chat and quizzes in a few moments.", 
            "filename": sanitized_filename
        }

    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(500, "Error processing document")


@router.get("/list/{course_id}")
async def list_documents(
    course_id: str,
    current_user: User = Depends(get_authenticated_user)
):
    """List all documents for a course"""
    try:
        user_id = str(current_user.id)
        from utils.file_utils import get_user_course_dir

        user_dir = get_user_course_dir(user_id, course_id)

        if not os.path.exists(user_dir):
            return {"documents": []}

        files = [
            f for f in os.listdir(user_dir)
            if not f.endswith(".annotations.json")
        ]
        return {"documents": files}
    except ValueError:
        raise HTTPException(400, "Invalid course id")
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(500, "Error listing documents")


@router.get("/paragraphs/{course_id}/{filename}")
async def get_paragraphs(
    course_id: str,
    filename: str,
    current_user: User = Depends(get_authenticated_user)
):
    """Return document split into indexed blocks + existing annotations"""
    user_id = str(current_user.id)
    file_path = _get_file_path(user_id, course_id, filename)

    if not os.path.exists(file_path):
        raise HTTPException(404, "Document not found")

    try:
        ext = os.path.splitext(filename)[1].lower()
        if ext in [".xlsx", ".xls", ".csv"]:
            import pandas as pd
            if ext == ".csv":
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Limit preview to 150 rows to ensure snappy UI performance
            df_preview = df.head(150)
            columns = [str(col) for col in df_preview.columns.tolist()]
            rows = df_preview.fillna("").astype(str).values.tolist()
            
            # Keep fallback text paragraph
            paragraphs = [df.to_string(index=False)]
            styles = ["body"]
            annotations = _load_annotations(_get_annotations_path(user_id, course_id, filename))
            
            return {
                "paragraphs": paragraphs,
                "styles": styles,
                "annotations": annotations,
                "data_table": {
                    "columns": columns,
                    "rows": rows,
                    "total_rows": len(df)
                }
            }

        blocks = _extract_paragraphs(file_path)
        # paragraphs list = text only (for backward compat with annotation index)
        paragraphs = [b["text"] for b in blocks]
        styles = [b["style"] for b in blocks]
        annotations = _load_annotations(_get_annotations_path(user_id, course_id, filename))
        return {"paragraphs": paragraphs, "styles": styles, "annotations": annotations}
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.error(f"Error reading document paragraphs: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(500, "Could not read document")


class AnnotationRequest(BaseModel):
    paragraph_index: int  # -1 means general / not tied to a paragraph
    page_index: Optional[int] = -1 # -1 means not tied to a page
    content: str
    annotation_type: Optional[str] = "note"


@router.get("/annotations/{course_id}/{filename}")
async def get_annotations(
    course_id: str,
    filename: str,
    current_user: User = Depends(get_authenticated_user)
):
    user_id = str(current_user.id)
    try:
        annotations = _load_annotations(_get_annotations_path(user_id, course_id, filename))
    except ValueError as e:
        raise HTTPException(400, str(e))
    return {"annotations": annotations}


@router.post("/annotations/{course_id}/{filename}")
async def save_annotation(
    course_id: str,
    filename: str,
    request: AnnotationRequest,
    current_user: User = Depends(get_authenticated_user)
):
    """Save an inline annotation tied to a paragraph index — deduplicates by content+paragraph+type"""
    user_id = str(current_user.id)
    try:
        ann_path = _get_annotations_path(user_id, course_id, filename)
    except ValueError as e:
        raise HTTPException(400, str(e))
    annotations = _load_annotations(ann_path)

    # Dedup: skip if identical content+paragraph+page+type already exists
    for existing in annotations:
        if (
            existing.get("paragraph_index") == request.paragraph_index
            and existing.get("page_index") == request.page_index
            and existing.get("content", "").strip() == request.content.strip()
            and existing.get("type") == request.annotation_type
        ):
            return {"message": "Duplicate annotation skipped", "total": len(annotations)}

    annotations.append({
        "id": int(datetime.utcnow().timestamp() * 1000),
        "paragraph_index": request.paragraph_index,
        "page_index": request.page_index,
        "content": request.content,
        "type": request.annotation_type,
        "created_at": datetime.utcnow().isoformat()
    })

    _save_annotations(ann_path, annotations)
    return {"message": "Annotation saved", "total": len(annotations)}


@router.delete("/annotations/{course_id}/{filename}/{annotation_id}")
async def delete_annotation(
    course_id: str,
    filename: str,
    annotation_id: int,
    current_user: User = Depends(get_authenticated_user)
):
    user_id = str(current_user.id)
    try:
        ann_path = _get_annotations_path(user_id, course_id, filename)
    except ValueError as e:
        raise HTTPException(400, str(e))

    if not os.path.exists(ann_path):
        raise HTTPException(404, "No annotations found")

    annotations = _load_annotations(ann_path)
    annotations = [a for a in annotations if a["id"] != annotation_id]
    _save_annotations(ann_path, annotations)
    return {"message": "Annotation deleted", "total": len(annotations)}


@router.get("/raw/{course_id}/{filename}")
async def serve_raw_file(
    course_id: str,
    filename: str,
    current_user: User = Depends(get_authenticated_user)
):
    """Stream the raw file so the frontend can embed it in an iframe."""
    user_id = str(current_user.id)
    try:
        file_path = _get_file_path(user_id, course_id, filename)
    except ValueError as e:
        raise HTTPException(400, str(e))

    if not os.path.exists(file_path):
        raise HTTPException(404, "Document not found")

    ext = os.path.splitext(filename)[1].lower()
    mime_map = {
        ".pdf":  "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".csv":  "text/csv",
        ".xml":  "application/xml",
        ".jpg":  "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png":  "image/png",
        ".txt":  "text/plain; charset=utf-8",
    }
    media_type = mime_map.get(ext, "application/octet-stream")

    def iter_file():
        with open(file_path, "rb") as f:
            while chunk := f.read(65536):
                yield chunk

    return StreamingResponse(
        iter_file(),
        media_type=media_type,
        headers={
            "Content-Disposition": f'inline; filename="{filename}"',
            "Cache-Control": "no-store",
        }
    )


@router.get("/export/{course_id}/{filename}")
async def export_annotated_document(
    course_id: str,
    filename: str,
    current_user: User = Depends(get_authenticated_user)
):
    """Export the annotated document as a Word (.docx) file"""
    user_id = str(current_user.id)
    try:
        file_path = _get_file_path(user_id, course_id, filename)
    except ValueError as e:
        raise HTTPException(400, str(e))

    if not os.path.exists(file_path):
        raise HTTPException(404, "Document not found")

    try:
        blocks = _extract_paragraphs(file_path)
        paragraphs = [b["text"] for b in blocks]
        annotations = _load_annotations(_get_annotations_path(user_id, course_id, filename))
        docx_bytes = _build_annotated_docx(paragraphs, annotations, filename)

        export_name = f"annotated_{os.path.splitext(filename)[0]}.docx"
        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{export_name}"'}
        )
    except Exception as e:
        logger.error(f"Export failed: {e}")
        logger.error(traceback.format_exc())
        raise HTTPException(500, "Export failed")



@lru_cache(maxsize=60)
def _render_page_cached(file_path: str, page_num: int, mtime: float) -> bytes:
    """Render a PDF page to PNG. mtime is included so cache invalidates if file changes."""
    doc = fitz.open(file_path)
    page = doc[page_num]
    mat = fitz.Matrix(2.0, 2.0)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    png_bytes = pix.tobytes("png")
    doc.close()
    return png_bytes


@router.get("/page/{course_id}/{filename}/{page_num}")
async def get_pdf_page_image(
    course_id: str,
    filename: str,
    page_num: int,
    current_user: User = Depends(get_authenticated_user)
):
    """Render a single PDF page as a PNG image using PyMuPDF (cached)."""
    user_id = str(current_user.id)
    try:
        file_path = _get_file_path(user_id, course_id, filename)
    except ValueError as e:
        raise HTTPException(400, str(e))

    if not os.path.exists(file_path):
        raise HTTPException(404, "Document not found")

    try:
        mtime = os.path.getmtime(file_path)
        # Validate page range first
        doc = fitz.open(file_path)
        total = len(doc)
        doc.close()
        if page_num < 0 or page_num >= total:
            raise HTTPException(404, f"Page {page_num} out of range (0-{total-1})")

        png_bytes = _render_page_cached(file_path, page_num, mtime)
        return StreamingResponse(
            io.BytesIO(png_bytes),
            media_type="image/png",
            headers={"Cache-Control": "max-age=3600"}
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Page render failed: {e}")
        raise HTTPException(500, "Page render failed")


@router.get("/pagecount/{course_id}/{filename}")
async def get_pdf_page_count(
    course_id: str,
    filename: str,
    current_user: User = Depends(get_authenticated_user)
):
    """Return total page count for a PDF."""
    user_id = str(current_user.id)
    try:
        file_path = _get_file_path(user_id, course_id, filename)
    except ValueError as e:
        raise HTTPException(400, str(e))

    if not os.path.exists(file_path):
        raise HTTPException(404, "Document not found")

    try:
        doc = fitz.open(file_path)
        count = len(doc)
        doc.close()
        return {"page_count": count}
    except Exception as e:
        logger.error(f"Page count failed: {e}")
        raise HTTPException(500, "Page count failed")


@router.get("/thumbnail/{course_id}")
async def get_course_thumbnail(
    course_id: str,
    current_user: User = Depends(get_authenticated_user)
):
    """Get or generate a thumbnail for a course based on its first document."""
    try:
        user_id = str(current_user.id)
        thumb_path = thumbnail_service.get_course_thumbnail(user_id, course_id)
        
        if thumb_path and os.path.exists(thumb_path):
            return FileResponse(thumb_path)
            
        # If no thumbnail could be generated, raise 404
        # The frontend can then use a fallback
        raise HTTPException(404, "No thumbnail available")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Thumbnail error: {e}")
        raise HTTPException(500, str(e))
